#!/usr/bin/env python3
"""Generate Strategic Review OnePager as .docx"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)
font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.25

ACCENT = RGBColor(0x7A, 0x7A, 0xE6)
GREEN = RGBColor(0x2a, 0x7d, 0x4f)
AMBER = RGBColor(0x9a, 0x6f, 0x1e)
RED = RGBColor(0xb5, 0x3d, 0x3d)
MUTED = RGBColor(0x5c, 0x5c, 0x5c)
DARK = RGBColor(0x1a, 0x1a, 0x1a)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = ACCENT if level == 2 else DARK
        run.font.name = 'Calibri'
    h.paragraph_format.space_before = Pt(14) if level == 1 else Pt(10)
    h.paragraph_format.space_after = Pt(6)
    return h

def add_para(text, bold=False, color=None, size=None, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return p

def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear'
    })
    shading.append(shd)

def style_cell(cell, text, bold=False, color=None, size=9, bg=None):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if bg:
        set_cell_shading(cell, bg)

# ══════════════════════════════════
# HEADER
# ══════════════════════════════════
p = doc.add_paragraph()
run = p.add_run('transform')
run.font.size = Pt(13)
run.font.color.rgb = DARK
run.font.name = 'Calibri'
run2 = p.add_run('az')
run2.font.size = Pt(13)
run2.font.color.rgb = ACCENT
run2.bold = True
run2.font.name = 'Calibri'

h = doc.add_heading('Proposal Review — SME Group', level=1)
for run in h.runs:
    run.font.color.rgb = DARK
    run.font.name = 'Calibri'
h.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
run = p.add_run('Agustín → Fermín  ·  Marzo 2026')
run.font.size = Pt(9)
run.font.color.rgb = MUTED
run.font.name = 'Calibri'

# ── Line ──
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(8)
run = p.add_run('─' * 80)
run.font.size = Pt(6)
run.font.color.rgb = RGBColor(0xe0, 0xdd, 0xd6)

# ══════════════════════════════════
# INTRO
# ══════════════════════════════════
p = doc.add_paragraph()
run = p.add_run('Fermín — revisé la propuesta que le mandaste a SME Group y te paso lo que encontré. En resumen: ')
run.font.size = Pt(9.5)
run.font.color.rgb = MUTED
run2 = p.add_run('el trabajo de fondo es sólido, pero le estás dejando plata en la mesa por cómo lo presentás.')
run2.font.size = Pt(9.5)
run2.bold = True
run2.font.color.rgb = DARK
run3 = p.add_run(' Abajo te dejo los números, lo que está bien, lo que cambiaría, y por qué.')
run3.font.size = Pt(9.5)
run3.font.color.rgb = MUTED
p.paragraph_format.space_after = Pt(10)

# ══════════════════════════════════
# CONTEXT TABLE
# ══════════════════════════════════
table = doc.add_table(rows=2, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Lo que cotizaste', 'Rate que mostrás', 'Lo que cobra el mercado']
values = ['£9,000', '£150/hr', '£15K–30K']
notes = ['2 fases · 60 hrs · 6 semanas', 'Visible en la propuesta', 'Por el mismo scope en UK']

for i, (h, v, n) in enumerate(zip(headers, values, notes)):
    cell = table.cell(0, i)
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(h.upper())
    run.font.size = Pt(7)
    run.font.color.rgb = MUTED
    run.bold = True
    run.font.name = 'Calibri'
    set_cell_shading(cell, 'FAF9F7')

    cell2 = table.cell(1, i)
    cell2.text = ''
    p2 = cell2.paragraphs[0]
    run2 = p2.add_run(v)
    run2.font.size = Pt(14)
    run2.bold = True
    run2.font.name = 'Calibri'
    run2.font.color.rgb = DARK
    p2.add_run('\n')
    run3 = p2.add_run(n)
    run3.font.size = Pt(7.5)
    run3.font.color.rgb = MUTED
    run3.font.name = 'Calibri'
    set_cell_shading(cell2, 'FAF9F7')

# Remove table borders
for row in table.rows:
    for cell in row.cells:
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.makeelement(qn('w:tcBorders'), {})
        for edge in ['top', 'left', 'bottom', 'right']:
            el = tcBorders.makeelement(qn(f'w:{edge}'), {
                qn('w:val'): 'none', qn('w:sz'): '0', qn('w:space'): '0'
            })
            tcBorders.append(el)
        tcPr.append(tcBorders)

doc.add_paragraph()  # spacer

# ══════════════════════════════════
# FINDINGS
# ══════════════════════════════════
add_heading_styled('Lo que encontré', level=2)

findings = [
    ('BIEN', GREEN, 'El rate está bien posicionado',
     '£150/hr queda en el rango alto de mid-level y bajo de senior specialist. Para boutique consulting en UK está correcto.'),
    ('PROBLEMA', RED, 'Pero lo estás mostrando',
     'Si el cliente ve horas × tarifa, compara contra freelancers a £50-80/hr. La pregunta pasa de "vale la pena resolver esto?" a "vale esta persona £150 la hora?"'),
    ('OJO', AMBER, 'Estás cobrando poco',
     '£9K es lo que cobran solo por el discovery. Diagnostic + implementation en UK va de £15K a £30K. Estás regalando la ejecución.'),
    ('FALTA', RED, 'No cuantificás el problema',
     'Explicás qué está mal pero nunca decís cuánto les cuesta. Sin un número en £/año, los £9K se juzgan en el aire.'),
    ('BIEN', GREEN, 'Entendés el negocio',
     'La sección de business understanding es lo mejor de la propuesta. Se nota que entendés el modelo de franquicia y los dos áreas de transformación.'),
    ('FLOJO', AMBER, 'No te protegés',
     '40 horas para 2 áreas de transformación es muy justo. No hay cláusula de ajuste de scope, ni de demoras del cliente, ni mecanismo de pivot.'),
    ('FALTA', RED, 'Cero proof points',
     'Ningún caso de éxito previo. Un solo dato concreto ("redujimos el ciclo un 65%") cambia completamente la credibilidad.'),
    ('FALTA', AMBER, 'Un solo precio',
     'Precio único = decisión sí/no. Con 3 opciones el cliente elige cuál, no si. El 73% elige la del medio y el deal size sube 20-30%.'),
]

for badge, badge_color, title, desc in findings:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)

    run_badge = p.add_run(f'[{badge}]  ')
    run_badge.font.size = Pt(7)
    run_badge.bold = True
    run_badge.font.color.rgb = badge_color
    run_badge.font.name = 'Calibri'

    run_title = p.add_run(title + '  —  ')
    run_title.font.size = Pt(9.5)
    run_title.bold = True
    run_title.font.color.rgb = DARK
    run_title.font.name = 'Calibri'

    run_desc = p.add_run(desc)
    run_desc.font.size = Pt(9)
    run_desc.font.color.rgb = MUTED
    run_desc.font.name = 'Calibri'

# ══════════════════════════════════
# RECOMMENDATIONS
# ══════════════════════════════════
add_heading_styled('Qué cambiaría (por orden de impacto)', level=2)

recs = [
    ('Sacá el rate por hora de la propuesta',
     'Mostrá un fee por proyecto. El detalle de horas va en el SOW si lo piden. Cambia toda la conversación.'),
    ('Ponele un número al problema del cliente',
     '"Esto les cuesta £X por año." Los £9K pasan a ser el 10-20% de lo que ahorran. ROI de 5-10x.'),
    ('Agregá aunque sea un caso previo',
     'Un solo dato concreto de experiencia anterior sube la credibilidad de forma desproporcionada.'),
    ('Armá 3 opciones de precio',
     'Básico / Recomendado / Premium. El cliente deja de pensar "contrato o no" y pasa a "cuál elijo".'),
    ('Sumá términos de pago y cláusulas de scope',
     'Te protege de scope creep, demoras del cliente y cambios de dirección.'),
    ('Subí a £12-15K o usá los £9K como opción base',
     'Estás por debajo del piso del mercado. Si metés una opción premium arriba, los £9K se ven "razonables".'),
]

table2 = doc.add_table(rows=len(recs) + 1, cols=3)
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
for i, h in enumerate(['#', 'Acción', 'Por qué']):
    cell = table2.cell(0, i)
    style_cell(cell, h, bold=True, color=MUTED, size=7.5)
    set_cell_shading(cell, 'F4F3F0')

# Data rows
for idx, (action, why) in enumerate(recs):
    row = idx + 1
    style_cell(table2.cell(row, 0), str(row), bold=True, color=ACCENT, size=9)
    style_cell(table2.cell(row, 1), action, bold=True, size=9)
    style_cell(table2.cell(row, 2), why, color=MUTED, size=8.5)

# Set column widths
table2.columns[0].width = Cm(0.8)
table2.columns[1].width = Cm(7)
table2.columns[2].width = Cm(8.5)

doc.add_paragraph()  # spacer

# ══════════════════════════════════
# TAKEAWAY
# ══════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
set_cell_bg = None  # We'll use a table for the dark box

takeaway_table = doc.add_table(rows=1, cols=1)
cell = takeaway_table.cell(0, 0)
set_cell_shading(cell, '1A1A1A')
cell.text = ''
p = cell.paragraphs[0]

run_label = p.add_run('EN UNA LÍNEA\n')
run_label.font.size = Pt(7)
run_label.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run_label.bold = True
run_label.font.name = 'Calibri'

run1 = p.add_run('El laburo que hacés es bueno y se nota que ')
run1.font.size = Pt(9.5)
run1.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
run1.font.name = 'Calibri'

run2 = p.add_run('entendés al cliente')
run2.font.size = Pt(9.5)
run2.font.color.rgb = RGBColor(0xeb, 0xeb, 0xfa)
run2.bold = True
run2.font.name = 'Calibri'

run3 = p.add_run('. El tema no es lo que ofrecés sino ')
run3.font.size = Pt(9.5)
run3.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
run3.font.name = 'Calibri'

run4 = p.add_run('cómo lo empaquetás')
run4.font.size = Pt(9.5)
run4.font.color.rgb = RGBColor(0xeb, 0xeb, 0xfa)
run4.bold = True
run4.font.name = 'Calibri'

run5 = p.add_run('. Con los mismos deliverables pero mejor framing comercial, esta propuesta deja de ser competitiva y pasa a ser difícil de rechazar.')
run5.font.size = Pt(9.5)
run5.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
run5.font.name = 'Calibri'

# Remove takeaway table borders
for row in takeaway_table.rows:
    for cell in row.cells:
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.makeelement(qn('w:tcBorders'), {})
        for edge in ['top', 'left', 'bottom', 'right']:
            el = tcBorders.makeelement(qn(f'w:{edge}'), {
                qn('w:val'): 'none', qn('w:sz'): '0', qn('w:space'): '0'
            })
            tcBorders.append(el)
        tcPr.append(tcBorders)

# ── Footer ──
doc.add_paragraph()
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
run = p.add_run('TransformAZ · Interno  |  Análisis completo: Proposal-Strategic-Review.html  |  Agustín · Marzo 2026')
run.font.size = Pt(7.5)
run.font.color.rgb = MUTED
run.font.name = 'Calibri'
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── Save ──
output_path = os.path.join(os.path.dirname(__file__), 'Strategic-Review-OnePager.docx')
doc.save(output_path)
print(f'DOCX generated: {output_path}')
